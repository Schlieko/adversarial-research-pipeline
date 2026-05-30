import os
import json
from dotenv import load_dotenv
from google import genai
from google.genai import types
from openai import OpenAI
from docx import Document
from docx.shared import Pt, RGBColor

# ==========================================
# 1. SETUP & AUTHENTICATION
# ==========================================
target_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(target_dir, ".env")) 

# Gemini stays native for deep-research compatibility
gemini_client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

# OpenRouter handles BOTH DeepSeek and Qwen
openrouter_client = OpenAI(
    api_key=os.getenv("OPENROUTER_API_KEY"), 
    base_url="https://openrouter.ai/api/v1"
)

# ==========================================
# 2. FILE MANAGEMENT (Models & Prompts)
# ==========================================
print("=== WELCOME TO THE 3-WAY COLLABORATIVE AI THINK-TANK ===")

# Dynamic models configuration
models_filepath = os.path.join(target_dir, "models.json")
if not os.path.exists(models_filepath):
    default_models = {
        "senior": [
            "gemini-3.1-pro-preview", 
            "gemini-2.5-pro"
        ],
        "junior": [
            "deepseek/deepseek-v4-pro", 
            "meta-llama/llama-3.3-70b-instruct:free",
            "deepseek/deepseek-v4-flash:free",
            "qwen/qwen3.7-max", 
            "nvidia/nemotron-3-super-120b-a12b:free",
            "moonshotai/kimi-k2.6:free"
        ],
        "grounding": [
            "qwen/qwen3.7-max", 
            "nvidia/nemotron-3-super-120b-a12b:free",
            "moonshotai/kimi-k2.6:free",
            "deepseek/deepseek-v4-pro", 
            "meta-llama/llama-3.3-70b-instruct:free",
            "deepseek/deepseek-v4-flash:free"
        ]
    }
    with open(models_filepath, "w", encoding="utf-8") as f:
        json.dump(default_models, f, indent=4)
    print("📁 Created clean 'models.json'.")

with open(models_filepath, "r", encoding="utf-8") as f:
    available_models = json.load(f)

# Base Prompt setup
seed_filepath = os.path.join(target_dir, "seed_prompt.md")
if not os.path.exists(seed_filepath):
    with open(seed_filepath, "w", encoding="utf-8") as f:
        f.write("Write your starting prompt or question here.")
    print("⚠️ Created 'seed_prompt.md'. Please write your starting prompt and run again.")
    exit()

# Attachments Setup
attachments_dir = os.path.join(target_dir, "prompt_files")
if not os.path.exists(attachments_dir):
    os.makedirs(attachments_dir)

with open(seed_filepath, "r", encoding="utf-8") as f:
    base_prompt = f.read().strip()

if base_prompt == "Write your starting prompt or question here." or not base_prompt:
    print("❌ Please update 'seed_prompt.md' with your actual question before running.")
    exit()

# Extract attachments text
attached_text = ""
for filename in os.listdir(attachments_dir):
    filepath = os.path.join(attachments_dir, filename)
    if filename.endswith((".txt", ".md")):
        with open(filepath, "r", encoding="utf-8") as file:
            attached_text += f"\n\n--- START ATTACHED FILE: {filename} ---\n{file.read()}\n--- END ATTACHED FILE ---\n"
    elif filename.endswith(".docx"):
        try:
            doc = Document(filepath)
            doc_text = "\n".join([para.text for para in doc.paragraphs])
            attached_text += f"\n\n--- START ATTACHED FILE: {filename} ---\n{doc_text}\n--- END ATTACHED FILE ---\n"
        except Exception:
            pass

# Save the original context as an anchor for the Grounding AI
original_context = base_prompt + attached_text
current_input = original_context

# ==========================================
# 3. TERMINAL UI (Model Selection)
# ==========================================
def select_model(role, model_list):
    print(f"\n=== SELECT YOUR {role.upper()} ===")
    for idx, model in enumerate(model_list):
        print(f"{idx + 1}. {model}")
    
    if role == "anchor ai (optional)":
        print("0. Skip (Run standard 2-way conversation)")
        
    choice = input("> ")
    if role == "anchor ai (optional)" and (choice == "0" or choice == ""):
        return None
    try:
        return model_list[int(choice) - 1]
    except:
        return model_list[0]

senior_model = select_model("first voice ai", available_models["senior"])
junior_model = select_model("second voice ai", available_models["junior"])
bs_model = select_model("anchor ai (optional)", available_models["grounding"])

rounds_input = input("\nHow many rounds of discussion would you like? (e.g., 5): ")
total_rounds = int(rounds_input) if rounds_input.isdigit() else 5

# ==========================================
# 4. INITIALIZE PERSONAS & MEMORY
# ==========================================
personas_filepath = os.path.join(target_dir, "personas.json")

# Auto-create personas.json if it doesn't exist using the new collaborative design
if not os.path.exists(personas_filepath):
    default_personas = {
        "senior": "Engage in a collaborative, lateral-thinking discussion. Analyze the user's prompt and files, thinking outside the box. Provide deep, expansive insights. Always end your turn by asking a specific, thought-provoking question for the next AI to answer.",
        "junior": "You are the second voice in a lateral-thinking AI discussion. Respond directly to the previous AI's output and answer their closing question. Add your own out-of-the-box ideas and expand the conversation in new directions. Always end your turn by asking a new question for the group to consider.",
        "grounding": "You are the anchor of this collaborative AI discussion. Review what the other AIs have discussed, then go all the way back to the user's original seed prompt and files. Tell the group what you think the human is really asking or trying to achieve. Add your own substantive, lateral thoughts to the conversation. Point out key things the group needs to consider moving forward to stay aligned with the human's core intent. End your turn by asking a guiding question for the next round."
    }
    with open(personas_filepath, "w", encoding="utf-8") as f:
        json.dump(default_personas, f, indent=4)
    print("📁 Created clean 'personas.json'.")

with open(personas_filepath, "r", encoding="utf-8") as f:
    personas = json.load(f)

gemini_chat = gemini_client.chats.create(
    model=senior_model, 
    config=types.GenerateContentConfig(system_instruction=personas["senior"])
)

junior_messages = [{"role": "system", "content": personas["junior"]}]
bs_messages = [{"role": "system", "content": personas["grounding"]}] if bs_model else []

junior_messages.append({"role": "user", "content": f"The starting topic is provided below.\n\n{current_input}"})

# ==========================================
# 5. THE COLLABORATIVE THINK-TANK LOOP
# ==========================================
markdown_log = f"# AI Think-Tank Log\n**Topic:** {base_prompt[:100]}...\n**Rounds:** {total_rounds}\n---\n\n"
with open("debate_log.md", "w", encoding="utf-8") as f:
    f.write(markdown_log)

print("\n[Starting the Collaborative Loop...]\n")

for i in range(total_rounds):
    print(f"--- Round {i+1} of {total_rounds} ---")
    
    # 1. GEMINI (First Voice - Native SDK)
    print("AI 1 (Gemini) is thinking...")
    gemini_response = gemini_chat.send_message(current_input)
    gemini_text = gemini_response.text
    
    with open("debate_log.md", "a", encoding="utf-8") as f:
        f.write(f"### AI Voice 1 ({senior_model}) - Round {i+1}\n{gemini_text}\n\n---\n\n")
    markdown_log += f"### AI Voice 1 - Round {i+1}\n{gemini_text}\n\n"

    # 2. DEEPSEEK/OTHER (Second Voice - Via OpenRouter)
    print("AI 2 is processing and expanding...")
    junior_messages.append({"role": "user", "content": gemini_text})
    
    ds_response = openrouter_client.chat.completions.create(
        model=junior_model,
        messages=junior_messages
    )
    ds_text = ds_response.choices[0].message.content
    junior_messages.append({"role": "assistant", "content": ds_text})
    
    with open("debate_log.md", "a", encoding="utf-8") as f:
        f.write(f"### AI Voice 2 ({junior_model}) - Round {i+1}\n{ds_text}\n\n---\n\n")
    markdown_log += f"### AI Voice 2 - Round {i+1}\n{ds_text}\n\n"

    current_input = ds_text

    # 3. QWEN/OTHER (Anchor AI - Via OpenRouter)
    if bs_model:
        print("Anchor AI is anchoring and synthesizing...")
        
        bs_payload = f"""
        ORIGINAL HUMAN INPUT & FILES:
        {original_context}
        
        ---
        AI 1 SAID:
        {gemini_text}
        
        AI 2 SAID:
        {ds_text}
        
        ---
        INSTRUCTIONS: Synthesize the conversation. Remind the group what the human is actually trying to achieve based on the original input. Add your own insights, and ask a guiding question for the next round.
        """
        bs_messages.append({"role": "user", "content": bs_payload})
        
        qwen_response = openrouter_client.chat.completions.create(
            model=bs_model,
            messages=bs_messages
        )
        qwen_text = qwen_response.choices[0].message.content
        bs_messages.append({"role": "assistant", "content": qwen_text})
        
        with open("debate_log.md", "a", encoding="utf-8") as f:
            f.write(f"### Anchor AI ({bs_model}) - Round {i+1}\n{qwen_text}\n\n---\n\n")
        markdown_log += f"### Anchor AI - Round {i+1}\n{qwen_text}\n\n"
        
        # Pass the collaborative baton back to the first AI
        current_input = f"AI 2's THOUGHTS: {ds_text}\n\nANCHOR AI's SYNTHESIS: {qwen_text}\n\nContinue the discussion, answer their questions, and explore further."

# ==========================================
# 6. EXECUTIVE SUMMARY (Word Doc Generation)
# ==========================================
print("\n[Discussion Complete! Generating Executive Summary Word Document...]")

summary_prompt = f"""
You are an executive summarizer. Review the preceding AI discussion transcript. 
Identify the core insights, strategic directions, and the most profound 
open questions or execution pathways raised. 

Format your response using ONLY clean Markdown headings (## ) and bullet points (* ). 
DO NOT use asterisks for bolding (**). Write in clear, professional plain text.

DISCUSSION TRANSCRIPT:
{markdown_log}
"""

try:
    summary_response = gemini_client.models.generate_content(
        model=senior_model,
        contents=summary_prompt
    )
    
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    
    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(17, 85, 204)
    
    doc.add_heading('AI Think-Tank: Executive Summary', 0)
    
    for line in summary_response.text.split('\n'):
        clean_line = line.replace('**', '').strip()
        if clean_line.startswith('## '):
            doc.add_heading(clean_line.replace('## ', ''), level=1)
        elif clean_line.startswith('### '):
            doc.add_heading(clean_line.replace('### ', ''), level=2)
        elif clean_line.startswith('* ') or clean_line.startswith('- '):
            doc.add_paragraph(clean_line[2:], style='List Bullet')
        elif clean_line:
            doc.add_paragraph(clean_line)
            
    doc.save("Executive_Summary.docx")
    print("✅ Success! Executive summary saved to 'Executive_Summary.docx'.")

except Exception as e:
    print(f"⚠️ Error generating summary: {e}")

print("\n=== SCRIPT 1 COMPLETE ===")
