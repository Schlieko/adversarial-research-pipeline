import os
import time
import json
from docx import Document
from docx.shared import Pt, RGBColor
from google import genai
from dotenv import load_dotenv

# ==========================================
# 1. SETUP & CONFIGURATION
# ==========================================
target_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(target_dir, ".env"))

api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    print("❌ ERROR: GEMINI_API_KEY not found in .env file.")
    exit()

client = genai.Client(api_key=api_key)

# Using the heavy reasoning model for maximum depth
model_id = 'gemini-3.1-pro-preview'

def main():
    print("=== STEP 2: PLANNER/WRITER DEEP RESEARCH LOOP ===")
    
    # ==========================================
    # 2. LOAD THE DEBATE LOG
    # ==========================================
    debate_filepath = os.path.join(target_dir, "debate_log.md")
    try:
        with open(debate_filepath, 'r', encoding='utf-8') as file:
            debate_content = file.read().strip()
            print("📂 Successfully loaded 'debate_log.md'.")
    except FileNotFoundError:
        print("❌ Error: 'debate_log.md' not found. Please run your debate script first.")
        return

    # ==========================================
    # 3. MANUAL FALLBACK PROMPT (For the Web UI)
    # ==========================================
    print("\n📝 Step 1: Generating Manual Web UI Fallback Prompt...")
    meta_prompt = f"""
    You are an expert Prompt Engineer. Based on the provided AI debate transcript, 
    write a highly optimized prompt that the user can copy and paste into the Gemini Advanced Deep Research UI.
    
    The prompt MUST instruct the AI to:
    1. Resolve the core disagreements highlighted by the Senior and Junior AIs.
    2. Investigate the specific hidden vulnerabilities or opportunities uncovered.
    3. Write a massive, highly detailed report (approx 20 pages) structured with clear headings.
    
    Output ONLY the text of the prompt.
    """
    try:
        prompt_response = client.models.generate_content(model=model_id, contents=meta_prompt)
        
        fallback_filepath = os.path.join(target_dir, "Manual_Deep_Research_Prompt.txt")
        with open(fallback_filepath, 'w', encoding='utf-8') as f:
            f.write("--- COPY & PASTE INTO GEMINI ADVANCED (DEEP RESEARCH) ---\n")
            f.write("--- DON'T FORGET TO ATTACH YOUR DEBATE_LOG.MD FILE ---\n\n")
            f.write(prompt_response.text.strip())
        print("✅ Saved fallback prompt to 'Manual_Deep_Research_Prompt.txt'")
    except Exception as e:
        print(f"⚠️ Error generating fallback prompt: {e}")

    # ==========================================
    # 4. GENERATE THE OUTLINE (The Planner)
    # ==========================================
    print("\n🧠 Step 2: Generating Comprehensive Report Outline...")
    outline_prompt = f"""
    You are an expert Research Director. Based on the provided multi-agent AI debate, 
    create a highly detailed, 8-section outline for a comprehensive Deep Research Report.
    
    The report must explore the initial topic, the points of friction between the AIs, and the ultimate conclusions reached.
    
    OUTPUT FORMAT:
    You must output ONLY valid JSON in the following format:
    [
      {{"section_title": "1. Executive Summary", "focus_instructions": "Summarize the core debate..."}},
      {{"section_title": "2. Primary Vulnerabilities", "focus_instructions": "Detail the flaws uncovered..."}}
    ]

    [START DEBATE TRANSCRIPT]
    {debate_content}
    [END DEBATE TRANSCRIPT]
    """
    try:
        outline_response = client.models.generate_content(model=model_id, contents=outline_prompt)
        raw_json = outline_response.text.strip().replace("```json", "").replace("```", "")
        outline = json.loads(raw_json)
        print(f"✅ Successfully generated an {len(outline)}-section outline.")
    except Exception as e:
        print(f"❌ Error generating outline: {e}. Check if the model output raw text instead of JSON.")
        return

    # ==========================================
    # 5. PREPARE THE WORD DOCUMENT
    # ==========================================
    print("\n🎨 Applying professional Word styling...")
    doc = Document()
    
    # Clean professional styling
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    
    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(17, 85, 204)
    
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(17, 85, 204)
    
    doc.add_heading('Deep Research Final Report', 0)

    # ==========================================
    # 6. GENERATE SECTIONS (The Writer Loop)
    # ==========================================
    print("\n✍️ Step 3: Generating Report Section by Section (This will take a few minutes)...")
    
    for index, section in enumerate(outline):
        title = section.get("section_title", f"Section {index + 1}")
        instructions = section.get("focus_instructions", "")
        
        print(f"   -> Drafting: {title}...")
        
        section_prompt = f"""
        You are an elite corporate researcher and Technical Writer. Write a highly detailed, comprehensive section for a Deep Research Report.
        Expand deeply on the concepts. Aim for 1,000 to 1,500 words for this section alone.
        
        CURRENT SECTION TITLE: {title}
        YOUR INSTRUCTIONS FOR THIS SECTION: {instructions}
        
        CRITICAL FORMATTING INSTRUCTIONS:
        1. DO NOT use raw Markdown links.
        2. DO NOT use asterisks (**) for bolding. Write in clean, professional plain text.
        3. Break up large walls of text into readable paragraphs.
        
        Use the context below to inform your writing.
        [START DEBATE TRANSCRIPT]
        {debate_content}
        [END DEBATE TRANSCRIPT]
        
        Output ONLY the text for this specific section. Do not include introductory filler.
        """

        try:
            section_response = client.models.generate_content(model=model_id, contents=section_prompt)
            doc.add_heading(title, level=1)
            doc.add_paragraph(section_response.text.strip())
            
            # Brief pause to respect API rate limits
            time.sleep(3) 
            
        except Exception as e:
            print(f"⚠️ Error drafting '{title}': {e}")
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"[ERROR GENERATING SECTION: {e}]")

    output_filepath = os.path.join(target_dir, "Final_Deep_Research_Report.docx")
    doc.save(output_filepath)
    print(f"\n🎉 Success! Massive report saved as 'Final_Deep_Research_Report.docx'")

    # ==========================================
    # 7. NOTEBOOKLM AUDIO PROMPT
    # ==========================================
    print("\n🎙️ Step 4: Generating NotebookLM Custom Instructions...")
    notebook_system_prompt = f"""
    You are an expert podcast producer. Based on the debate transcript below, write a highly optimized 'Custom Instruction' for Google NotebookLM's Audio Overview feature.

    Tell the two AI hosts:
    1. The primary thesis of the debate.
    2. The specific tone they should adopt (e.g., investigative, enthusiastic, cautious).
    3. 2-3 specific controversies or insights from the debate they absolutely must highlight.

    Output ONLY the instruction text (under 2000 characters).

    DEBATE TRANSCRIPT:
    {debate_content[:6000]} # Using the first chunk of the debate for context
    """

    try:
        notebook_response = client.models.generate_content(model=model_id, contents=notebook_system_prompt)
        
        instructions_filepath = os.path.join(target_dir, "NotebookLM_Custom_Instructions.txt")
        with open(instructions_filepath, "w", encoding="utf-8") as f:
            f.write("--- COPY & PASTE INTO NOTEBOOKLM'S AUDIO OVERVIEW CUSTOM INSTRUCTIONS ---\n\n")
            f.write(notebook_response.text.strip())
            
        print("✅ Saved NotebookLM instructions to 'NotebookLM_Custom_Instructions.txt'")
    except Exception as e:
        print(f"⚠️ Error generating NotebookLM instructions: {e}")

    print("\n=== PIPELINE COMPLETE! ===")

if __name__ == "__main__":
    main()